# from docx import Document
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.styles import getSampleStyleSheet

# # WORD FILE
# def generate_word(ddr_data, file_path):
#     doc = Document()

#     doc.add_heading("Detailed Diagnostic Report", 0)

#     doc.add_heading("Summary", 1)
#     doc.add_paragraph(ddr_data.get("summary", "Not Available"))

#     doc.add_heading("Observations", 1)
#     observations = ddr_data.get("observations", [])
#     for obs in observations:
#         doc.add_paragraph(f"- {obs}")

#     doc.add_heading("Root Cause", 1)
#     doc.add_paragraph(ddr_data.get("root_cause", "Not Available"))

#     doc.add_heading("Severity", 1)
#     doc.add_paragraph(ddr_data.get("severity", "Not Available"))

#     doc.add_heading("Recommendations", 1)
#     doc.add_paragraph(ddr_data.get("recommendations", "Not Available"))

#     doc.add_heading("Missing Information", 1)
#     doc.add_paragraph(ddr_data.get("missing_info", "Not Available"))

#     doc.save(file_path)


# # PDF FILE
# def generate_pdf(ddr_data, file_path):
#     doc = SimpleDocTemplate(file_path)
#     styles = getSampleStyleSheet()

#     content = []

#     content.append(Paragraph("Detailed Diagnostic Report", styles["Title"]))
#     content.append(Spacer(1, 10))

#     content.append(Paragraph(ddr_data.get("summary", ""), styles["Normal"]))
#     content.append(Spacer(1, 10))

#     content.append(Paragraph(ddr_data.get("root_cause", ""), styles["Normal"]))
#     content.append(Spacer(1, 10))

#     content.append(Paragraph(ddr_data.get("severity", ""), styles["Normal"]))
#     content.append(Spacer(1, 10))

#     content.append(Paragraph(ddr_data.get("recommendations", ""), styles["Normal"]))

#     doc.build(content)






from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


# ===================== WORD FILE =====================
def generate_word(ddr_data, file_path):
    doc = Document()

    doc.add_heading("Detailed Diagnostic Report", 0)

    # Summary
    doc.add_heading("Summary", 1)
    doc.add_paragraph(ddr_data.get("summary", "Not Available"))

    # Observations
    doc.add_heading("Observations", 1)
    observations = ddr_data.get("observations", [])
    if observations:
        for obs in observations:
            doc.add_paragraph(f"• {obs}")
    else:
        doc.add_paragraph("Not Available")

    # Root Cause
    doc.add_heading("Root Cause", 1)
    doc.add_paragraph(ddr_data.get("root_cause", "Not Available"))

    # Severity
    doc.add_heading("Severity", 1)
    doc.add_paragraph(ddr_data.get("severity", "Not Available"))

    # Recommendations
    doc.add_heading("Recommendations", 1)
    doc.add_paragraph(ddr_data.get("recommendations", "Not Available"))

    # Missing Info
    doc.add_heading("Missing Information", 1)
    doc.add_paragraph(ddr_data.get("missing_info", "Not Available"))

    doc.save(file_path)


# ===================== PDF FILE =====================
def generate_pdf(ddr_data, file_path):
    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()

    content = []

    # Helper function
    def section(title, value):
        content.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        content.append(Spacer(1, 8))
        content.append(Paragraph(value, styles["Normal"]))
        content.append(Spacer(1, 12))

    # Title
    content.append(Paragraph("Detailed Diagnostic Report", styles["Title"]))
    content.append(Spacer(1, 20))

    # Summary
    section("Summary", ddr_data.get("summary", "Not Available"))

    # Observations
    observations = ddr_data.get("observations", [])
    if observations:
        obs_text = "<br/>".join([f"• {o}" for o in observations])
    else:
        obs_text = "Not Available"
    section("Observations", obs_text)

    # Root Cause
    section("Root Cause", ddr_data.get("root_cause", "Not Available"))

    # Severity
    section("Severity", ddr_data.get("severity", "Not Available"))

    # Recommendations
    section("Recommendations", ddr_data.get("recommendations", "Not Available"))

    # Missing Info
    section("Missing Information", ddr_data.get("missing_info", "Not Available"))

    doc.build(content)